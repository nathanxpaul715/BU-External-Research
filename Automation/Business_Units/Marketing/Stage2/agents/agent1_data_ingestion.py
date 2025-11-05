"""Agent 1: Data Ingestion & Context Builder
Loads all input files and structures them for downstream agents
"""
import pandas as pd
from docx import Document
from typing import Dict, List, Any
import os
import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import (
    BU_INTELLIGENCE_PATH,
    USE_CASES_CSV_PATH,
    FUNCTION_UPDATES_CSV_PATH,
    OPTIONAL_FILES
)


class DataIngestionAgent:
    """Agent responsible for loading and structuring all input data"""

    def __init__(self):
        self.bu_intelligence = None
        self.use_cases = None
        self.function_updates = None
        self.optional_data = {}
        self.context = {}

    def load_bu_intelligence(self) -> str:
        """Load BU Intelligence document and extract full text"""
        print("Loading BU Intelligence document...")
        doc = Document(BU_INTELLIGENCE_PATH)

        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Also extract tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    full_text.append(row_text)

        self.bu_intelligence = "\n\n".join(full_text)
        print(f"[OK] Loaded BU Intelligence: {len(self.bu_intelligence)} characters")
        return self.bu_intelligence

    def load_use_cases(self) -> pd.DataFrame:
        """Load AI Use Cases CSV"""
        print("Loading AI Use Cases CSV...")
        self.use_cases = pd.read_csv(USE_CASES_CSV_PATH)
        print(f"[OK] Loaded {len(self.use_cases)} use cases")
        print(f"  Columns: {list(self.use_cases.columns)}")
        return self.use_cases

    def load_function_updates(self) -> pd.DataFrame:
        """Load Function Updates CSV"""
        print("Loading Function Updates CSV...")
        self.function_updates = pd.read_csv(FUNCTION_UPDATES_CSV_PATH)
        print(f"[OK] Loaded {len(self.function_updates)} function updates")
        return self.function_updates

    def load_optional_files(self) -> Dict[str, str]:
        """Load optional files if they exist"""
        print("Checking for optional files...")
        for key, path in OPTIONAL_FILES.items():
            if os.path.exists(path):
                try:
                    doc = Document(path)
                    text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    self.optional_data[key] = text
                    print(f"[OK] Loaded optional file: {key}")
                except Exception as e:
                    print(f"[FAIL] Failed to load {key}: {e}")
            else:
                print(f"[SKIP] Optional file not found: {key}")
        return self.optional_data

    def extract_key_context(self) -> Dict[str, Any]:
        """Extract key context elements from BU Intelligence for quick reference"""
        print("Extracting key context...")

        # This is a simplified extraction - in production, you might use LLM to extract structured data
        context = {
            "full_bu_intelligence": self.bu_intelligence,
            "bu_intelligence_length": len(self.bu_intelligence),
            "num_use_cases": len(self.use_cases) if self.use_cases is not None else 0,
            "use_case_list": self.use_cases["Use Case Name"].tolist() if self.use_cases is not None else [],
            "optional_files_loaded": list(self.optional_data.keys())
        }

        # Extract some key sections (if identifiable)
        bu_text = self.bu_intelligence.upper()
        context["has_competitors_section"] = "COMPETITOR" in bu_text
        context["has_vendors_section"] = "VENDOR" in bu_text or "TECHNOLOGY" in bu_text
        context["has_strategy_section"] = "STRATEGY" in bu_text or "STRATEGIC" in bu_text

        self.context = context
        print(f"[OK] Context extracted: {len(context)} elements")
        return context

    def prepare_use_cases_for_enrichment(self) -> List[Dict[str, Any]]:
        """Prepare use cases as structured dictionaries for enrichment"""
        print("Preparing use cases for enrichment...")

        use_cases_list = []
        for idx, row in self.use_cases.iterrows():
            use_case = {
                "index": idx,
                "function": row.get("Segment / Function", "Marketing"),
                "original_name": row.get("Use Case Name", ""),
                "original_description": row.get("Use Case Description", "") if pd.notna(row.get("Use Case Description")) else "",
                "original_outcomes": row.get("Outcomes/Deliverables", "") if pd.notna(row.get("Outcomes/Deliverables")) else "",
                "ai_tools": row.get("AI Tools", "") if pd.notna(row.get("AI Tools")) else "",
                "stage": row.get("Stage", "") if pd.notna(row.get("Stage")) else "",
                "strategy": row.get("Use Case Strategy", "") if pd.notna(row.get("Use Case Strategy")) else "",
                "full_row": row.to_dict()
            }
            use_cases_list.append(use_case)

        print(f"[OK] Prepared {len(use_cases_list)} use cases for enrichment")
        return use_cases_list

    def run(self) -> Dict[str, Any]:
        """Execute all data ingestion tasks"""
        print("=" * 80)
        print("AGENT 1: DATA INGESTION & CONTEXT BUILDER")
        print("=" * 80)

        # Load all files
        self.load_bu_intelligence()
        self.load_use_cases()
        self.load_function_updates()
        self.load_optional_files()

        # Extract context
        self.extract_key_context()

        # Prepare use cases
        use_cases_prepared = self.prepare_use_cases_for_enrichment()

        result = {
            "bu_intelligence": self.bu_intelligence,
            "use_cases": use_cases_prepared,
            "function_updates": self.function_updates,
            "optional_data": self.optional_data,
            "context": self.context
        }

        print("\n[OK] DATA INGESTION COMPLETE")
        print("=" * 80)
        return result


if __name__ == "__main__":
    # Test the agent
    agent = DataIngestionAgent()
    result = agent.run()
    print(f"\nResult keys: {list(result.keys())}")
    print(f"Number of use cases: {len(result['use_cases'])}")
