# agents/report_writer_agent.py
import os
from datetime import datetime
from docx import Document
from config import settings

class ReportWriterAgent:
    """
    Stage 4 AI-Native Reimagination Blueprint Writer.
    Takes prioritized use cases and writes a structured DOCX blueprint
    with specific sections and tables as per business requirements.
    """

    def __init__(self, prioritized_records: list, output_basename: str = "stage4_blueprint"):
        self.records = prioritized_records or []
        self.output_basename = output_basename

    def run(self) -> str:
        if not self.records:
            raise ValueError("No prioritized records provided to ReportWriterAgent.")

        print("[ReportWriterAgent] Generating AI-Native Reimagination Blueprint DOCX...")

        doc = Document()

        # Cover Page
        doc.add_heading("AI-Native Reimagination Blueprint", level=0)
        doc.add_paragraph(f"Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph("Business Function: [Specify here in orchestrator params]")
        doc.add_paragraph("This blueprint presents prioritized AI/automation use cases for strategic consideration.")

        doc.add_page_break()

        # Intro Section
        doc.add_heading("1. Executive Summary", level=1)
        doc.add_paragraph(
            "This document synthesizes internal and external intelligence to identify and prioritize "
            "AI/automation opportunities. The use cases are ranked by business impact, complexity, "
            "confidence, and strategic fit."
        )

        # Table of Use Cases
        doc.add_heading("2. Prioritized Use Cases", level=1)
        table = doc.add_table(rows=1, cols=9)
        headers = [
            "Priority Rank", "Title", "Description", "Category",
            "Impact", "Complexity", "Confidence", "Strategic Fit", "Type"
        ]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for rank, rec in enumerate(self.records, start=1):
            row = table.add_row().cells
            row[0].text = str(rank)
            row[1].text = rec.get("title", "")
            row[2].text = rec.get("description", "")
            row[3].text = rec.get("category", "")
            row[4].text = str(rec.get("impact", ""))
            row[5].text = str(rec.get("complexity", ""))
            row[6].text = rec.get("confidence", "")
            row[7].text = str(rec.get("strategic_fit", ""))
            row[8].text = rec.get("type", "")

        # Additional strategic recommendations section
        doc.add_page_break()
        doc.add_heading("3. Strategic Recommendations", level=1)
        doc.add_paragraph(
            "Based on the analysis, the following actions are recommended:\n"
            "- Focus on top-ranked use cases for immediate implementation.\n"
            "- Allocate exploratory budget for non-core opportunities with high impact potential.\n"
            "- Initiate feasibility assessments for medium complexity items."
        )

        # Save file
        output_path = os.path.join(settings.OUTPUT_DIR, f"{self.output_basename}.docx")
        doc.save(output_path)

        print(f"[ReportWriterAgent] Blueprint saved to {output_path}")
        return output_path