# utils/file_parser.py  
import pandas as pd  
import docx  
from typing import Dict, List, Optional, Tuple  
from loguru import logger  
import openpyxl
  
class FileParser:  
    """Handles parsing of all input files for the agentic system"""
      
    def __init__(self):  
        logger.info("FileParser initialized")
      
    def parse_bu_intelligence(self, file_path: str) -> Dict[str, str]:  
        """  
        Parse BU Intelligence Document (Word/PDF)  
        Returns structured business context  
        """  
        logger.info(f"Parsing BU Intelligence document: {file_path}")
          
        try:  
            # Parse Word document  
            doc = docx.Document(file_path)  
            full_text = []
              
            for paragraph in doc.paragraphs:  
                if paragraph.text.strip():  
                    full_text.append(paragraph.text)
              
            content = "\n".join(full_text)
              
            # Extract key sections  
            sections = {  
                "executive_summary": self._extract_section(content, "EXECUTIVE SUMMARY", "STRATEGIC FOUNDATION"),  
                "strategic_foundation": self._extract_section(content, "STRATEGIC FOUNDATION", "ORGANIZATIONAL DESIGN"),  
                "organizational_design": self._extract_section(content, "ORGANIZATIONAL DESIGN", "OPERATIONAL ARCHITECTURE"),  
                "operational_architecture": self._extract_section(content, "OPERATIONAL ARCHITECTURE", "TECHNOLOGY ECOSYSTEM"),  
                "technology_ecosystem": self._extract_section(content, "TECHNOLOGY ECOSYSTEM", "VENDOR & SERVICE"),  
                "full_content": content  
            }
              
            logger.success(f"Successfully parsed BU Intelligence document with {len(content)} characters")  
            return sections
              
        except Exception as e:  
            logger.error(f"Failed to parse BU Intelligence document: {e}")  
            raise
      
    def parse_enriched_use_cases(self, file_path: str) -> pd.DataFrame:  
        """  
        Parse Enriched Use Cases Excel file  
        Returns DataFrame with existing use cases for deduplication  
        """  
        logger.info(f"Parsing Enriched Use Cases: {file_path}")
          
        try:  
            # Read the Enriched Use Cases sheet  
            df = pd.read_excel(file_path, sheet_name="Enriched Use Cases")
              
            logger.info(f"Found {len(df)} existing use cases for deduplication")  
            logger.success(f"Successfully parsed Enriched Use Cases: {len(df)} rows")
              
            return df
              
        except Exception as e:  
            logger.error(f"Failed to parse Enriched Use Cases: {e}")  
            raise
      
    def parse_role_activity_mapping(self, file_path: str) -> Dict[str, pd.DataFrame]:  
        """  
        Parse Role Activity Mapping Excel file  
        Extracts ALL sub-functions (sheets) for autonomous processing  
        Returns dict with sheet_name: DataFrame  
        """  
        logger.info(f"Parsing Role Activity Mapping Excel: {file_path}")
          
        try:  
            # Load Excel file and get all sheet names  
            excel_file = pd.ExcelFile(file_path)  
            all_sheets = excel_file.sheet_names
              
            # Exclude Summary sheet  
            sub_function_sheets = [s for s in all_sheets if s.lower() != "summary"]
              
            logger.info(f"Found {len(sub_function_sheets)} sub-function sheets: {sub_function_sheets}")
              
            # Parse each sub-function sheet  
            sub_functions = {}  
            for sheet_name in sub_function_sheets:  
                df = pd.read_excel(file_path, sheet_name=sheet_name)  
                sub_functions[sheet_name] = df  
                logger.info(f"  - Parsed '{sheet_name}': {len(df)} rows, {len(df.columns)} columns")
              
            logger.success(f"Successfully parsed {len(sub_functions)} sub-functions from Role Activity Mapping")
              
            return sub_functions
              
        except Exception as e:  
            logger.error(f"Failed to parse Role Activity Mapping: {e}")  
            raise
      
    def parse_strategic_priorities(self, file_path: str) -> Dict[str, any]:  
        """  
        Parse Strategic Priority documents (PDF/Word)  
        Returns structured priority data  
        """  
        logger.info(f"Parsing Strategic Priorities document: {file_path}")
          
        try:  
            # Parse Word document  
            doc = docx.Document(file_path)  
            full_text = []
              
            for paragraph in doc.paragraphs:  
                if paragraph.text.strip():  
                    full_text.append(paragraph.text)
              
            content = "\n".join(full_text)
              
            priorities = {  
                "full_content": content,  
                "file_name": file_path.split("/")[-1]  
            }
              
            logger.success(f"Successfully parsed Strategic Priorities document")  
            return priorities
              
        except Exception as e:  
            logger.error(f"Failed to parse Strategic Priorities: {e}")  
            raise
      
    def _extract_section(self, content: str, start_marker: str, end_marker: str) -> str:  
        """Helper to extract section between markers"""  
        try:  
            start_idx = content.find(start_marker)  
            end_idx = content.find(end_marker)
              
            if start_idx != -1 and end_idx != -1:  
                return content[start_idx:end_idx].strip()  
            elif start_idx != -1:  
                return content[start_idx:].strip()  
            else:  
                return ""  
        except:  
            return ""
      
    def extract_sub_function_metadata(self, df: pd.DataFrame, sheet_name: str) -> Dict[str, any]:  
        """  
        Extract metadata from a sub-function DataFrame  
        Returns structured metadata for processing  
        """  
        try:  
            # Extract activities  
            activities = []
              
            for idx, row in df.iterrows():  
                activity = row.get("Activity", "")  
                time_spent = row.get("Estimated % of Time Spent", 0)  
                fte = row.get("FTE", 0)  
                functional_tools = row.get("Functional AI Tools (Applicable)", "")  
                generic_tools = row.get("Generic AI Assistants (Applicable)", "")  
                ai_usage = row.get("% usage of existing AI tools for activity", 0)  
                existing_use_cases = row.get("Existing AI Use Cases\n(List out the activities the existing AI tools are being used for here)", "")  
                proposed_use_cases = row.get("Proposed Use Cases\n(AI use cases identified but not deployed, wishlist)", "")
                  
                if pd.notna(activity) and activity.strip() and activity.lower() not in ["total", "note"]:  
                    activities.append({  
                        "activity": activity,  
                        "time_spent_pct": float(time_spent) if pd.notna(time_spent) else 0,  
                        "fte": float(fte) if pd.notna(fte) else 0,  
                        "functional_tools": str(functional_tools) if pd.notna(functional_tools) else "",  
                        "generic_tools": str(generic_tools) if pd.notna(generic_tools) else "",  
                        "ai_usage_pct": float(ai_usage) if pd.notna(ai_usage) else 0,  
                        "existing_use_cases": str(existing_use_cases) if pd.notna(existing_use_cases) else "",  
                        "proposed_use_cases": str(proposed_use_cases) if pd.notna(proposed_use_cases) else ""  
                    })
              
            metadata = {  
                "sheet_name": sheet_name,  
                "total_activities": len(activities),  
                "activities": activities  
            }
              
            logger.info(f"Extracted metadata for '{sheet_name}': {len(activities)} activities")  
            return metadata
              
        except Exception as e:  
            logger.error(f"Failed to extract metadata from {sheet_name}: {e}")  
            raise